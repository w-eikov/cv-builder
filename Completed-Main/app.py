from flask import Flask, request, render_template, send_file
import pdfkit
import tempfile
import os
import openai
import json
import re
from datetime import datetime
from docx import Document

app = Flask(__name__)

# ---------------- PDF CONFIG ----------------
WKHTMLTOPDF_PATH = r"C:/Program Files/wkhtmltopdf/bin/wkhtmltopdf.exe"
pdf_config = pdfkit.configuration(wkhtmltopdf=WKHTMLTOPDF_PATH)

# ---------------- AI CONFIG ----------------
openai.api_key = "" #THIS KEY MUST NOT BE SHARED BY ANY MEANS, so it won't be on the github lmao

# ---------------- Helpers ----------------
def extract_json_from_text(text: str) -> str:
    if not text:
        return text
    m = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, flags=re.DOTALL | re.IGNORECASE)
    if m:
        return m.group(1)
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and start < end:
        return text[start:end + 1]
    return text


def clean_greeting(g: str) -> str:
    if not g:
        return ""
    g = g.strip()
    g = re.sub(r"^\s*dear\s*[:,]?\s*", "", g, flags=re.IGNORECASE)
    g = g.rstrip(",")
    return g.strip()


def nl_to_br(s: str) -> str:
    return s.replace("\n", "<br>") if s else s


def paragraphs_to_html(body: str) -> str:
    if not body:
        return ""
    paras = re.split(r"\n\s*\n", body.strip())
    return "".join(f"<p>{p.strip()}</p>" for p in paras if p.strip())


def ensure_keys(obj: dict, keys: list[str]) -> dict:
    return {k: (obj.get(k, "") if obj.get(k, "") is not None else "") for k in keys}


# ---------------- AI: Improve & Split ----------------
def improve_and_split_cover_letter(raw_text: str) -> dict:
    prompt = (
        "You are an expert career coach and formatting assistant. "
        "Take the following raw cover letter and:\n"
        "1) Improve the writing to be polished, professional, clear, and concise.\n"
        "2) Return ONLY a valid JSON object with EXACTLY these keys:\n"
        '   - "client_info": 4 lines (Name, Address, Phone, Email).\n'
        '   - "recruiter_info": 5 lines with the FIRST line as a full date like "August 22, 2025", then Recruiter Name, Title, Bank, Address.\n'
        '   - "greeting": JUST the person’s name/title (e.g., "Ms. Smith"). Do NOT include "Dear" or commas.\n'
        '   - "body": 4 paragraphs separated by a blank line (two newlines between paragraphs).\n'
        '   - "closing": The closing with "Sincerely," on the first line, then the applicant’s full name on the next line.\n'
        "When referring generically to the organization in the body of the letter, always use the word firm exactly (do not substitute with bank, company, organization, etc.), even if the recruiter’s organization name contains such words. "
        "Strictly output a single JSON object only. Do not wrap in markdown/code fences."
    )

    response = openai.chat.completions.create(
        model="gpt-4.1",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": raw_text},
        ],
    )

    content = response.choices[0].message.content.strip()
    candidate = extract_json_from_text(content)

    try:
        data = json.loads(candidate)
    except json.JSONDecodeError:
        data = {
            "client_info": "",
            "recruiter_info": "",
            "greeting": "",
            "body": content,
            "closing": "",
        }

    data = ensure_keys(data, ["client_info", "recruiter_info", "greeting", "body", "closing"])

    # Normalize values so downstream formatters always receive strings
    for k, v in data.items():
        if isinstance(v, list):
            data[k] = "\n".join(str(item) for item in v)
        elif v is None:
            data[k] = ""

    data["greeting"] = clean_greeting(data.get("greeting", ""))
    return data


# ---------------- CV ROUTES ----------------
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/index.html")
def go_back_to_index():
    return render_template("/index.html")


@app.route("/form", methods=["GET", "POST"])
def form():
    if request.method == "POST":
        name = request.form.get("name")
        email = request.form.get("email")
        phone = request.form.get("phone")
        experiences = request.form.getlist("experience")
        skills = [s.strip() for s in request.form.get("skills", "").split(",") if s.strip()]
        goals = [s.strip() for s in request.form.get("goals", "").split(",") if s.strip()]
        interests = request.form.get("interests")

        return render_template(
            "cv_template.html",
            name=name,
            email=email,
            phone=phone,
            experiences=experiences,
            skills=skills,
            goals=goals,
            interests=interests,
        )
    return render_template("form.html")


@app.route("/preview-cv", methods=["POST"])
def preview_cv():
    name = request.form.get("name")
    email = request.form.get("email")
    phone = request.form.get("phone")
    experience_str = "\n".join(request.form.getlist("experience"))
    experiences = [e.strip() for e in experience_str.split("\n") if e.strip()]
    skills = [s.strip() for s in request.form.get("skills", "").split(",") if s.strip()]
    goals = [s.strip() for s in request.form.get("goals", "").split(",") if s.strip()]
    interests = request.form.get("interests")

    return render_template(
        "cv_template.html",
        name=name,
        email=email,
        phone=phone,
        experiences=experiences,
        skills=skills,
        goals=goals,
        interests=interests,
    )


@app.route("/download-cv", methods=["POST"])
def download_cv():
    name = request.form.get("name").replace(" ", "_")
    email = request.form.get("email")
    phone = request.form.get("phone")
    experience_str = request.form.get("experience", "")
    experiences = [e.strip() for e in experience_str.split("\n") if e.strip()]
    skills = [s.strip() for s in request.form.get("skills", "").split(",") if s.strip()]
    goals = [s.strip() for s in request.form.get("goals", "").split(",") if s.strip()]
    interests = request.form.get("interests")

    rendered = render_template(
        "cv_template.html",
        name=name,
        email=email,
        phone=phone,
        experiences=experiences,
        skills=skills,
        goals=goals,
        interests=interests,
    )

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmpfile:
        pdf_file_path = tmpfile.name

    pdfkit.from_string(
        rendered,
        pdf_file_path,
        configuration=pdf_config,
        options={"enable-local-file-access": ""},
    )

    response = send_file(pdf_file_path, as_attachment=True)

    @response.call_on_close
    def remove_file():
        os.remove(pdf_file_path)

    return response


# ---------------- COVER LETTER ROUTES ----------------
@app.route("/preview-coverletter", methods=["POST"])
def preview_coverletter():
    form_data = request.form.to_dict()
    raw_text = render_template("cover_letter_template.html", **form_data)
    return raw_text


@app.route("/download-coverletter", methods=["POST"])
def download_coverletter():
    form_data = request.form.to_dict()
    rendered = render_template("cover_letter_template.html", **form_data)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmpfile:
        pdf_file_path = tmpfile.name

    pdfkit.from_string(
        rendered,
        pdf_file_path,
        configuration=pdf_config,
        options={"enable-local-file-access": ""},
    )

    response = send_file(pdf_file_path, as_attachment=True)

    @response.call_on_close
    def remove_file():
        os.remove(pdf_file_path)

    return response


@app.route("/download-improved-coverletter", methods=["POST"])
def download_improved_coverletter():
    form_data = request.form.to_dict()
    raw_text = render_template("cover_letter_template.html", **form_data)
    sections = improve_and_split_cover_letter(raw_text)

    fallback_greeting = (
        form_data.get("cl_greeting")
        or form_data.get("cl_recruiter_name")
        or "Hiring Manager"
    )
    sections["greeting"] = clean_greeting(sections.get("greeting") or fallback_greeting)

    sections["client_info"] = nl_to_br(sections.get("client_info", "").strip())
    sections["recruiter_info"] = nl_to_br(sections.get("recruiter_info", "").strip())
    sections["closing"] = nl_to_br(sections.get("closing", "").strip())
    sections["body"] = paragraphs_to_html(sections.get("body", ""))

    rendered = render_template("ai_letter_template.html", **sections)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmpfile:
        pdf_file_path = tmpfile.name

    pdfkit.from_string(
        rendered,
        pdf_file_path,
        configuration=pdf_config,
        options={"enable-local-file-access": ""},
    )

    response = send_file(pdf_file_path, as_attachment=True)

    @response.call_on_close
    def remove_file():
        os.remove(pdf_file_path)

    return response


@app.route("/download-improved-coverletter-docx", methods=["POST"])
def download_improved_coverletter_docx():
    form_data = request.form.to_dict()
    raw_text = render_template("cover_letter_template.html", **form_data)
    sections = improve_and_split_cover_letter(raw_text)

    fallback_greeting = (
        form_data.get("cl_greeting")
        or form_data.get("cl_recruiter_name")
        or "Hiring Manager"
    )
    sections["greeting"] = clean_greeting(sections.get("greeting") or fallback_greeting)

    doc = Document()

    for line in sections.get("client_info", "").split("\n"):
        doc.add_paragraph(line, style="Normal").alignment = 2  # right

    doc.add_paragraph()

    for line in sections.get("recruiter_info", "").split("\n"):
        doc.add_paragraph(line, style="Normal").alignment = 0  # left

    doc.add_paragraph()
    doc.add_paragraph(f"Dear {sections.get('greeting', 'Hiring Manager')},")

    for para in sections.get("body", "").split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.add_paragraph()
    for line in sections.get("closing", "").split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmpfile:
        docx_file_path = tmpfile.name
        doc.save(docx_file_path)

    response = send_file(docx_file_path, as_attachment=True)

    @response.call_on_close
    def remove_file():
        os.remove(docx_file_path)

    return response


if __name__ == "__main__":
    app.run(debug=True)
